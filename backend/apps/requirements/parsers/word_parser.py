"""
Word document parser implementation.
"""
from typing import List, Dict
from docx import Document
from .base import BaseFileParser


class WordParser(BaseFileParser):
    """
    Parser for Microsoft Word documents (.docx).
    Extracts text from paragraphs and tables.
    """

    def parse(self, file_path: str) -> List[Dict]:
        """
        Parse Word document and extract data.

        Args:
            file_path: Path to Word document

        Returns:
            List of dictionaries containing extracted content
        """
        try:
            doc = Document(file_path)
            data = []

            # Process paragraphs
            for para in doc.paragraphs:
                text = para.text.strip()
                if text:
                    data.append({
                        'type': 'paragraph',
                        'text': text,
                        'style': para.style.name if para.style else 'Normal',
                    })

            # Process tables
            for table_idx, table in enumerate(doc.tables):
                # Extract headers from first row
                headers = []
                if table.rows:
                    headers = [cell.text.strip() for cell in table.rows[0].cells]

                # Extract data rows
                for row_idx, row in enumerate(table.rows[1:], 1):  # Skip header row
                    row_data = {}
                    for col_idx, cell in enumerate(row.cells):
                        key = headers[col_idx] if col_idx < len(headers) else f'column_{col_idx}'
                        row_data[key] = cell.text.strip()

                    # Only add if row has data
                    if any(v for v in row_data.values()):
                        row_data['type'] = 'table_row'
                        row_data['table_index'] = table_idx
                        row_data['row_index'] = row_idx
                        data.append(row_data)

            return data

        except Exception as e:
            raise ValueError(f"Failed to parse Word document: {str(e)}")

    def extract_requirements(self, parsed_data: List[Dict]) -> List[str]:
        """
        Extract requirements from Word document data.

        Args:
            parsed_data: List of dictionaries from parse()

        Returns:
            List of requirement text strings
        """
        requirements = []

        for item in parsed_data:
            # Extract from paragraphs
            if item.get('type') == 'paragraph':
                text = item.get('text', '').strip()
                if text:
                    # Skip very short texts (likely titles)
                    if len(text) > 10:
                        requirements.append(text)

            # Extract from tables
            elif item.get('type') == 'table_row':
                # Look for requirement-related fields
                requirement_text = None
                for key in ['requirement', 'feature', 'capability', 'description', 'text']:
                    if key in item and item[key] and item[key].strip():
                        requirement_text = item[key].strip()
                        break

                # If no standard field, use first substantial value
                if not requirement_text:
                    for value in item.values():
                        if isinstance(value, str) and len(value.strip()) > 10:
                            requirement_text = value.strip()
                            break

                if requirement_text:
                    requirements.append(requirement_text)

        return requirements

    def extract_structured_requirements(self, parsed_data: List[Dict]) -> List[Dict]:
        """
        Extract structured requirements with metadata.

        Args:
            parsed_data: List of dictionaries from parse()

        Returns:
            List of dictionaries with requirement text and metadata
        """
        structured_requirements = []

        for item in parsed_data:
            if item.get('type') == 'paragraph':
                text = item.get('text', '').strip()
                if text and len(text) > 10:
                    structured_requirements.append({
                        'text': text,
                        'source': 'paragraph',
                        'style': item.get('style', 'Normal'),
                    })

            elif item.get('type') == 'table_row':
                # Include all table data
                structured_requirements.append({
                    **item,
                    'source': 'table',
                })

        return structured_requirements
